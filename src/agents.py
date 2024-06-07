from autogen import AssistantAgent, ConversableAgent, UserProxyAgent
from autogen.coding import LocalCommandLineCodeExecutor
from autogen.cache import Cache
from autogen.agentchat.contrib.math_user_proxy_agent  import MathUserProxyAgent
from src.config import llm_config
from docx import Document
from docx.text.paragraph import Paragraph
import os
from typing import Optional, List, Tuple
from markdown2 import markdown
from bs4 import BeautifulSoup
from src.prompts import file_writer_agent_system_message, writer_system_message, critic_system_message, financial_reviewer_system_message, outliner_system_message, quality_system_message, layman_system_message, planner_system_message
import copy
import pprint
import re
from typing import Dict, List, Tuple

import autogen
from autogen.agentchat.contrib.capabilities import transform_messages, transforms


# -------- CODEX

executor = LocalCommandLineCodeExecutor(
    timeout=3600,
    work_dir="./src/codex",
)


file_writer_agent = ConversableAgent(
    name="file_writer_agent",
    system_message = file_writer_agent_system_message,
    llm_config=llm_config,
    code_execution_config={"executor": executor},
    human_input_mode="NEVER",
    default_auto_reply=
    "Please continue. If everything is done, reply 'TERMINATE'.",
)


# -------- Writer

writer = AssistantAgent(
    name="Writer",
    system_message=writer_system_message,
    llm_config=llm_config,
)

outliner = AssistantAgent(
    name="Writer",
    system_message=outliner_system_message,
    llm_config=llm_config,
)

# -------- Reviewers

critic = AssistantAgent(
    name="Critic",
    is_termination_msg=lambda x: x.get("content", "").find("TERMINATE") >= 0,
    llm_config=llm_config,
    system_message=critic_system_message,
)

layman_reviewer = AssistantAgent(
    name="Layman Reviewer",
    description="A reviewer that makes sure a laywoman would fully understand the content provided to her.",
    llm_config=llm_config,
    system_message=layman_system_message,
)

financial_reviewer = AssistantAgent(
    name="Financial Reviewer",
    description='A reviewer that makes sure financial justification are credible',
    llm_config=llm_config,
    system_message=financial_reviewer_system_message,
)

quality_reviewer = AssistantAgent(
    name="Quality Assurance Reviewer",
    description="a reviewer that makes sure that claims are well justified",
    llm_config=llm_config,
    system_message=quality_system_message,
)

meta_reviewer = AssistantAgent(
    name="Meta Reviewer",
    llm_config=llm_config,
    system_message="You are a meta reviewer, you aggragate and review "
    "the work of other reviewers and give a final suggestion on the content.",
)

# --------  Math


mathproxyagent = MathUserProxyAgent(
    name="mathproxyagent",
    human_input_mode="NEVER",
    code_execution_config={"use_docker": False},
)

# --------  Planner



planner = AssistantAgent(
    name="planner",
    llm_config=llm_config,
    # the default system message of the AssistantAgent is overwritten here
    system_message=planner_system_message,
)

planner_user = UserProxyAgent(
    name="planner_user",
    max_consecutive_auto_reply=0,  # terminate without auto-reply
    human_input_mode="NEVER",
    code_execution_config={
        "use_docker": False
    },  # Please set use_docker=True if docker is available to run the generated code. Using docker is safer than running the generated code directly.
)

def ask_planner(message):
    planner_user.initiate_chat(planner, message=message)
    # return the last message received from the planner
    return planner_user.last_message()["content"]

assistant = AssistantAgent(
    name="assistant",
    llm_config={
        "temperature": 0,
        "timeout": 600,
        "cache_seed": 42,
        "config_list": llm_config,
        "functions": [
            {
                "name": "ask_planner",
                "description": "ask planner to: 1. get a plan for finishing a task, 2. verify the execution result of the plan and potentially suggest new plan.",
                "parameters": {
                    "type": "object",
                    "properties": {
                        "message": {
                            "type": "string",
                            "description": "question to ask planner. Make sure the question include enough context, such as the code and the execution result. The planner does not know the conversation between you and the user, unless you share the conversation with the planner.",
                        },
                    },
                    "required": ["message"],
                },
            },
        ],
    },
)

# create a UserProxyAgent instance named "user_proxy"
user_proxy = UserProxyAgent(
    name="user_proxy",
    human_input_mode="TERMINATE",
    max_consecutive_auto_reply=10,
    # is_termination_msg=lambda x: "content" in x and x["content"] is not None and x["content"].rstrip().endswith("TERMINATE"),
    code_execution_config={
        "work_dir": "planning",
        "use_docker": False,
    },  # Please set use_docker=True if docker is available to run the generated code. Using docker is safer than running the generated code directly.
    function_map={"ask_planner": ask_planner},
)


# --------  Tools

@file_writer_agent.register_for_execution()
@file_writer_agent.register_for_llm(description="Writes a text string to a markdown file")
def write_text_to_markdown(text, file_name, directory='./src/result/intermediate_results'):
    """
    Writes a text string to a markdown file, ensuring the directory exists.

    Args:
        text (str): The raw text content to be written to the markdown file.
        file_name (str): The name of the markdown file.
        directory (str): The directory where the markdown file will be saved. Default is './output_folder'.
    """
    # Ensure the directory exists
    os.makedirs(directory, exist_ok=True)
    
    # Full path to the markdown file
    file_path = os.path.join(directory, file_name)
    
    # Write the text content to the markdown file
    with open(file_path, 'w', encoding='utf-8') as file:
        file.write(text)
    
    print(f'Text has been written to {file_path}')

@file_writer_agent.register_for_execution()
@file_writer_agent.register_for_llm(description="Parses a markdown string into several markdown strings divided by titles, then saves each text string as a separate markdown document in a folder.")
def parse_markdown(markdown_str, output_folder="./src/result/intermediate_results"):
    """
    Parses a markdown string into several markdown strings divided by titles,
    then saves each text string as a separate markdown document in a folder.

    Args:
        markdown_str (str): The markdown content as a string.
        output_folder (str): The folder to save the numbered markdown files.

    Returns:
        List of filenames of the created markdown documents.
    """
    # Regex pattern to match markdown titles (assuming titles start with #)
    title_pattern = re.compile(r'(#+\s.*\n)')
    
    # Find all titles and their start positions
    matches = list(title_pattern.finditer(markdown_str))
    
    # Ensure the output folder exists
    os.makedirs(output_folder, exist_ok=True)
    
    filenames = []
    
    # Iterate over the matches and extract content
    for i in range(len(matches)):
        start_pos = matches[i].start()
        end_pos = matches[i + 1].start() if i + 1 < len(matches) else len(markdown_str)
        
        section_str = markdown_str[start_pos:end_pos]
        
        # Write the section to a separate markdown file
        filename = os.path.join(output_folder, f'section_{i + 1}.md')
        with open(filename, 'w', encoding='utf-8') as f:
            f.write(section_str)
        
        filenames.append(filename)
    
    return filenames

@file_writer_agent.register_for_execution()
@file_writer_agent.register_for_llm(description="Reads a markdown file and returns its raw text content.")
def read_markdown_file_to_text(file_path):
    """
    Reads a markdown file and returns its raw text content.

    Args:
        file_path (str): The path to the markdown file.

    Returns:
        str: The raw text content extracted from the markdown file.
    """
    # Read the markdown file
    with open(file_path, 'r', encoding='utf-8') as file:
        markdown_string = file.read()

    # Convert markdown to HTML
    html = markdown.markdown(markdown_string)
    
    # Use BeautifulSoup to extract text from HTML
    soup = BeautifulSoup(html, 'html.parser')
    text = soup.get_text()
    
    return text

@file_writer_agent.register_for_execution()
@file_writer_agent.register_for_llm(description="Reads and loads the data from a .docx file as plain text and returns the complete file content as plain text.")
def read_docx_as_plain_text(filepath: str = '.src/result/result.docx') -> Optional[str]:
    """
    Reads and loads the data from a .docx file as plain text and returns the complete file content as plain text.

    :param filepath: The path to the .docx file to be read.
    :return: The plain text content of the .docx file, or None if the file cannot be read.
    """
    try:
        # Load the Document
        doc = Document(filepath)
        
        # Initialize an empty string to hold the text
        full_text = []
        
        # Iterate through each paragraph in the document
        for para in doc.paragraphs:
            full_text.append(para.text)
        
        # Join all paragraphs into a single string separated by newlines
        result_text = "\n".join(full_text)
        
        return result_text
    except Exception as e:
        print(f"An error occurred while reading the .docx file: {e}")
        return None
    
@file_writer_agent.register_for_execution()
@file_writer_agent.register_for_llm(description="Appends text to the end of a given title in a .docx file (before the next title) and saves the updated document.")
def append_text_to_title( title: str, text_to_append: str, doc_path: str= '.src/result/result.docx', save_path: Optional[str] = None ) -> None:
    """
    Converts markdown to .docx and appends text to the end of a given title in a .docx file (before the next title) and saves the updated document.

    :param title: The title to search for in the document.
    :param text_to_append: The text to append to the end of the section with the specified title.
    :return: None
    """
    try:
        # Convert markdown text to Docx-suitable format
        formatted_text = markdown_to_docx_format(text_to_append)
     
    try:
        # Load the Document
        doc = Document(doc_path)
        
        # Iterate through paragraphs to find the title
        found_title = False
        insertion_point = None

        for i, para in enumerate(doc.paragraphs):
            if found_title:
                if para.style.name.startswith('Heading'):
                    # We have reached the next title
                    insertion_point = i
                    break
            elif para.text.strip() == title:
                found_title = True

        # If the title is found, insert the text
        if found_title:
            if insertion_point is None:
                # Append to the end of the document if no next title is found
                doc.add_paragraph(formatted_text)
            else:
                # Insert before the next title
                doc.paragraphs.insert(insertion_point, Paragraph(text_to_append, doc))
        
        # Save the updated document
        if save_path is None:
            save_path = doc_path
        doc.save(save_path)
        print(f"Document updated and saved at: {save_path}")

    except Exception as e:
        print(f"An error occurred while processing the .docx file: {e}")

# # Example usage
# doc_path = ".src/result/result.docx"
# title_to_search = "Introduction"
# text_to_append = "This is the appended text."
# append_text_to_title(doc_path, title_to_search, text_to_append)

@file_writer_agent.register_for_execution()
@file_writer_agent.register_for_llm(description="Replaces the text between the specified title and the next title in a .docx file with new text.")
def replace_section_with_text(title: str, new_text: str, doc_path: str= '.src/result/result.docx',  save_path: Optional[str] = None) -> None:
    """
    Converts Markdown text and Replaces the text between the specified title and the next title in a .docx file with new text.

    :param doc_path: The path to the .docx file to be read.
    :param title: The title that marks the beginning of the section to replace.
    :param new_text: The new text to insert in place of the specified section.
    :param save_path: The path to save the updated .docx file. If None, overwrites the original file.
    :return: None
    """
    try:
        # Convert markdown text to Docx-suitable format
        formatted_text = markdown_to_docx_format(new_text)
     
    try:
        # Load the Document
        doc = Document(doc_path)
        
        # Initialize variables to track the section boundaries
        section_start = None
        section_end = None

        # Find the title and the boundaries of the section to replace
        for i, para in enumerate(doc.paragraphs):
            if section_start is not None and para.style.name.startswith('Heading'):
                section_end = i
                break
            if para.text.strip() == title:
                section_start = i

        # If the title is found, replace the section
        if section_start is not None:
            # Remove old section paragraphs
            if section_end is None:
                # If no next title, clear till the end of the document
                section_end = len(doc.paragraphs)
            for _ in range(section_end - section_start - 1):
                del doc.paragraphs[section_start + 1]
                
            # Insert new text
            new_paragraphs = formatted_text.split("\n")
            for new_para in new_paragraphs:
                doc.paragraphs[section_start].insert_paragraph_before(new_para)
            del doc.paragraphs[section_start]  # Remove the old title paragraph as the new ones include the title

        # Save the updated document
        if save_path is None:
            save_path = doc_path
        doc.save(save_path)
        print(f"Document updated and saved at: {save_path}")

    except Exception as e:
        print(f"An error occurred while processing the .docx file: {e}")

@file_writer_agent.register_for_execution()
@file_writer_agent.register_for_llm(description="Converts markdown content to a single string with python-docx compatible formatting.")
def markdown_to_docx_format(markdown_text: str) -> str:
    """
    Converts markdown content to a single string with python-docx compatible formatting.

    :param markdown_text: The raw text containing markdown elements.
    :return: A single string with the complete text, formatted for further processing.
    """
    try:
        # Convert markdown to HTML
        html = markdown(markdown_text)
        
        # Parse HTML using BeautifulSoup
        soup = BeautifulSoup(html, 'html.parser')
        
        formatted_text = ""

        # Define mapping for header tags
        header_map = {
            'h1': ('Heading 1', 1),
            'h2': ('Heading 2', 2),
            'h3': ('Heading 3', 3),
            'h4': ('Heading 4', 4),
            'h5': ('Heading 5', 5),
            'h6': ('Heading 6', 6)
        }
        
        # Helper function to add text with possible styling
        def add_styled_text(element):
            text = ''
            if element.name == 'strong':
                text += '**' + element.text + '**'
            elif element.name == 'em':
                text += '*' + element.text + '*'
            elif element.name == 'code':
                text += '`' + element.text + '`'
            else:
                text += element.text
            return text

        # Iterate through HTML elements
        for tag in soup.contents:
            if tag.name in header_map:
                formatted_text += f"\n\n{header_map[tag.name][0]}: {tag.text.strip()}\n"
            elif tag.name == 'p':
                p_text = ''.join(add_styled_text(child) for child in tag.children)
                formatted_text += f"\n{p_text}\n"
            elif tag.name == 'ul':
                for li in tag.find_all('li'):
                    formatted_text += f"\n- {li.text.strip()}"
            elif tag.name == 'ol':
                for li in tag.find_all('li'):
                    formatted_text += f"\n1. {li.text.strip()}"
            elif tag.name == 'pre':
                code_block = tag.find('code')
                if code_block:
                    formatted_text += f"\n\n```\n{code_block.text}\n```\n"

        return formatted_text.strip()


@file_writer_agent.register_for_execution()
@file_writer_agent.register_for_llm(description="Converts markdown text to docx format and writes it to a .docx file.")
def write_markdown_to_docx(markdown_text: str, save_path: Optional[str] = "./src/result/result.docx", doc_path: str = "./src/result/result.docx", ) -> None:
    """
    Converts markdown text to docx format and writes it to a .docx file.

    :param markdown_text: The raw text containing markdown elements.
    :param doc_path: The path to the existing .docx file.
    :param save_path: The path to save the updated .docx file. If None, overwrites the original file.
    :return: None
    """
    try:
        # Convert markdown text to Docx-suitable format
        formatted_text = markdown_to_docx_format(markdown_text)
        
        # Load the Document
        doc = Document(doc_path)
        
        # Split the formatted text into paragraphs based on newline separators
        paragraphs = formatted_text.split('\n\n')
        
        for paragraph in paragraphs:
            if ': ' in paragraph:
                # Check for headers
                header_type, text = paragraph.split(': ', 1)
                doc.add_heading(text.strip(), level=header_type[-1])
            elif paragraph.startswith('```') and paragraph.endswith('```'):
                # Check for code blocks
                code_text = paragraph.strip('```').strip()
                doc.add_paragraph(code_text, style='Quote')
            else:
                # Handle bullet points and other text
                lines = paragraph.split('\n')
                for line in lines:
                    if line.startswith('- '):
                        doc.add_paragraph(line.strip('- '), style='List Bullet')
                    elif line.startswith('1. '):
                        doc.add_paragraph(line.strip('1. '), style='List Number')
                    else:
                        doc.add_paragraph(line)

        # Save the updated document
        if save_path is None:
            save_path = doc_path
        doc.save(save_path)
        print(f"Document updated and saved at: {save_path}")

    except Exception as e:
        print(f"An error occurred while processing the .docx file: {e}")


def replace_specific_text_in_markdown(file_name: str, old_text: str, new_text: str , markdown_directory: str = "./src/result/intermediate_results"):
    """
    Opens a markdown file from a directory, replaces specified text, and saves/overrides the file.

    Args:
        markdown_directory (str): The directory containing the markdown file.
        file_name (str): The name of the markdown file to be modified.
        old_text (str): The text to search for and replace in the markdown file.
        new_text (str): The text to replace the old text with in the markdown file.
    
    Returns:
        None
    """
    # Construct the full file path
    file_path = os.path.join(markdown_directory, file_name)
    
    # Ensure the file exists in the specified directory
    if not os.path.isfile(file_path):
        raise ValueError(f"The specified markdown file does not exist: {file_path}")

    # Read the content of the markdown file
    with open(file_path, 'r', encoding='utf-8') as file:
        content = file.read()

    # Replace the specified text
    updated_content = content.replace(old_text, new_text)

    # Save the updated content back to the same file, overriding it
    with open(file_path, 'w', encoding='utf-8') as file:
        file.write(updated_content)

    print(f"Replaced text in {file_path} and saved the updated file.")

def replace_all_text_in_markdown(markdown_directory: str, file_name: str, new_text: str):
    """
    Opens a markdown file from a directory, replaces all text with new text, and saves/overrides the file.

    Args:
        markdown_directory (str): The directory containing the markdown file.
        file_name (str): The name of the markdown file to be modified.
        new_text (str): The new text to replace the existing content in the markdown file.
    
    Returns:
        None
    """
    # Construct the full file path
    file_path = os.path.join(markdown_directory, file_name)
    
    # Ensure the file exists in the specified directory
    if not os.path.isfile(file_path):
        raise ValueError(f"The specified markdown file does not exist: {file_path}")

    # Save the new content to the same file, overriding it
    with open(file_path, 'w', encoding='utf-8') as file:
        file.write(new_text)

    print(f"Replaced all text in {file_path} and saved the updated file.")

def replace_section_in_markdown(file_name: str, title: str, new_text: str, markdown_directory: str = ".src/result/intermediate_results"):
    """
    Opens a markdown file from a directory, finds a given title, and replaces all text
    until the next title or subtitle with the text provided, then saves/overrides the file.

    Args:
        markdown_directory (str): The directory containing the markdown file.
        file_name (str): The name of the markdown file to be modified.
        title (str): The title of the section to replace.
        new_text (str): The new text to replace the existing content in the section.
    
    Returns:
        None
    """
    # Construct the full file path
    file_path = os.path.join(markdown_directory, file_name)
    
    # Ensure the file exists in the specified directory
    if not os.path.isfile(file_path):
        raise ValueError(f"The specified markdown file does not exist: {file_path}")
    
    # Read the content of the markdown file
    with open(file_path, 'r', encoding='utf-8') as file:
        content = file.read()

    # Create a regex pattern to find the title and the subsequent content
    title_pattern = re.compile(r'^(#+\s+' + re.escape(title) + r'\s*\n)', re.MULTILINE)
    match = title_pattern.search(content)
    
    if not match:
        raise ValueError(f"Title '{title}' not found in the markdown file.")
    
    # Find the start and end positions of the section to replace
    start_pos = match.start()
    end_pos = len(content)
    
    next_title_pattern = re.compile(r'^#+\s+', re.MULTILINE)
    next_match = next_title_pattern.search(content, pos=match.end())
    
    if next_match:
        end_pos = next_match.start()
    
    # Create the new content by replacing the old section with the new text
    new_content = content[:match.end()] + new_text + content[end_pos:]
    
    # Save the updated content back to the same file, overriding it
    with open(file_path, 'w', encoding='utf-8') as file:
        file.write(new_content)

    print(f"Section under title '{title}' replaced and file saved at {file_path}.")

@file_writer_agent.register_for_execution()
@file_writer_agent.register_for_llm(description="Iterates over each markdown file in a directory and writes them to a .docx file.")
def process_markdown_files_in_directory(markdown_directory: str = "./src/result/intermediate_results", docx_save_path: str = "./src/result/result.docx"):
    """
    Iterates over each markdown file in a directory and writes them to a .docx file.

    Args:
        markdown_directory (str): The directory containing markdown files.
        docx_save_path (str): The path to save the combined .docx file.

    Returns:
        None
    """
    # Ensure the markdown directory exists
    if not os.path.isdir(markdown_directory):
        raise ValueError(f"The specified directory does not exist: {markdown_directory}")
    
    # Create a new Document or load an existing one
    if os.path.exists(docx_save_path):
        doc = Document(docx_save_path)
    else:
        doc = Document()
    
    # Iterate over each markdown file in the directory
    for filename in os.listdir(markdown_directory):
        if filename.endswith('.md'):
            filepath = os.path.join(markdown_directory, filename)
            
            # Read the markdown file content
            with open(filepath, 'r', encoding='utf-8') as file:
                markdown_text = file.read()
            
            # Convert markdown text to docx format and append it to the document
            formatted_text = markdown_to_docx_format(markdown_text)
            paragraphs = formatted_text.split('\n\n')
            
            for paragraph in paragraphs:
                if ': ' in paragraph:
                    # Check for headers
                    header_type, text = paragraph.split(': ', 1)
                    doc.add_heading(text.strip(), level=int(header_type[-1]))
                elif paragraph.startswith('```') and paragraph.endswith('```'):
                    # Check for code blocks
                    code_text = paragraph.strip('```').strip()
                    doc.add_paragraph(code_text, style='Quote')
                else:
                    # Handle bullet points and other text
                    lines = paragraph.split('\n')
                    for line in lines:
                        if line.startswith('- '):
                            doc.add_paragraph(line.strip('- '), style='List Bullet')
                        elif line.startswith('1. '):
                            doc.add_paragraph(line.strip('1. '), style='List Number')
                        else:
                            doc.add_paragraph(line)

    # Save the combined document
    doc.save(docx_save_path)
    print(f"Combined document saved at: {docx_save_path}")

# # Example usage
# doc_path = ".src/result/result.docx"
# title_to_search = "Introduction"
# new_section_text = """
# This is the new content to replace the section.
# You can add multiple paragraphs as needed.
# """
# replace_section_with_text(doc_path, title_to_search, new_section_text)


# # Simulated writing team function (generates content for each section)
# def writer_team(section_title):
#     # This function should generate content for a given section. 
#     # Here we just return a placeholder text.
#     return f"Content for {section_title}"

# def generate_document_from_outline(markdown_outline, output_folder="./src/result/intermediate_results", docx_save_path="./src/result/result.docx"):
#     os.makedirs(output_folder, exist_ok=True)
#     sections = parse_markdown(markdown_outline, output_folder)
#     for section_file in sections:
#         section_title = read_markdown_file_to_text(section_file)
#         generated_content = writer_team(section_title.strip())
#         write_text_to_markdown(generated_content, os.path.basename(section_file), output_folder)
#     process_markdown_files_in_directory(output_folder, docx_save_path)



# @file_writer_agent.register_for_execution()
# @file_writer_agent.register_for_llm(description="Converts markdown content to a .docx file with appropriate headers and text formatting.")
# def markdown_to_docx(markdown_text: str, doc_path: str) -> None:
#     """
#     Converts markdown content to a .docx file with appropriate headers and text formatting.

#     :param markdown_text: The raw text containing markdown elements.
#     :param doc_path: The path to save the generated .docx file.
#     :return: None
#     """
#     try:
#         # Convert markdown to HTML
#         html = markdown(markdown_text)
        
#         # Parse HTML using BeautifulSoup
#         soup = BeautifulSoup(html, 'html.parser')
        
#         # Create a new Document
#         doc = Document()
        
#         # Define mapping for header tags
#         header_map = {
#             'h1': 'Heading 1',
#             'h2': 'Heading 2',
#             'h3': 'Heading 3',
#             'h4': 'Heading 4',
#             'h5': 'Heading 5',
#             'h6': 'Heading 6'
#         }
        
#         # Helper function to add text with possible styling
#         def add_styled_text(p, element):
#             if element.name == 'strong':
#                 run = p.add_run(element.text)
#                 run.bold = True
#             elif element.name == 'em':
#                 run = p.add_run(element.text)
#                 run.italic = True
#             elif element.name == 'code':
#                 run = p.add_run(element.text)
#                 run.font.name = 'Courier New'
#             else:
#                 p.add_run(element.text)

#         # Iterate through HTML elements
#         for tag in soup.contents:
#             if tag.name in header_map:
#                 doc.add_heading(tag.text.strip(), level=int(tag.name[-1]))
#             elif tag.name == 'p':
#                 p = doc.add_paragraph()
#                 for child in tag.children:
#                     add_styled_text(p, child)
#             elif tag.name == 'ul':
#                 for li in tag.find_all('li'):
#                     doc.add_paragraph(li.text.strip(), style='List Bullet')
#             elif tag.name == 'ol':
#                 for li in tag.find_all('li'):
#                     doc.add_paragraph(li.text.strip(), style='List Number')
#             elif tag.name == 'pre':
#                 code_block = tag.find('code')
#                 if code_block:
#                     doc.add_paragraph(code_block.text, style='Code')

#         # Save the document
#         doc.save(doc_path)
#         print(f"Document saved at: {doc_path}")

#     except Exception as e:
#         print(f"An error occurred while processing the markdown text: {e}")


        
# @file_writer_agent.register_for_execution()
# @file_writer_agent.register_for_llm(description="create a docx from plain text with the right format for correctly rendering a docx")
# def create_docx_from_script(script: str, filename: str = 'result.docx', directory: str = '.src/result') -> None:
#     """
#     Creates a .docx file from a given script and saves it to the specified directory.

#     :param script: The text content to be saved in the .docx file.
#     :param filename: The name of the .docx file to be created. Default is 'result.docx'.
#     :param directory: The directory where the .docx file will be saved. Default is '.src/result'.
#     :return: None

#     Example usage
#     script_text = "
#     This is an example script that will be saved into the .docx file.
#     You can add more content here as needed.
#     "
#     create_docx_from_script(script_text)
#     """
#     # Create directory if it does not exist
#     if not os.path.exists(directory):
#         os.makedirs(directory)
    
#     # Create a new Document
#     doc = Document()
    
#     # Add the script to the document
#     doc.add_paragraph(script)
    
#     # Construct the full path
#     full_path = os.path.join(directory, filename)
    
#     # Save the document
#     doc.save(full_path)
#     print(f"Document saved at: {full_path}")


# code_executor_agent = ConversableAgent(
#     name="code_executor_agent",
#     llm_config=False,
#     code_execution_config={"executor": executor},
#     human_input_mode="ALWAYS",
#     default_auto_reply=
#     "Please continue. If everything is done, reply 'TERMINATE'.",
# )