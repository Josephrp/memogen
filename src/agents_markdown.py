from autogen import AssistantAgent, ConversableAgent, UserProxyAgent
from autogen.coding import LocalCommandLineCodeExecutor
from autogen.cache import Cache
from autogen.agentchat.contrib.math_user_proxy_agent  import MathUserProxyAgent
from src.config import llm_config
from docx import Document
from docx.text.paragraph import Paragraph
import os
from typing import Optional, List, Tuple
from markdown2 import markdown
from bs4 import BeautifulSoup
from src.prompts import file_writer_agent_system_message, writer_system_message, critic_system_message, financial_reviewer_system_message, outliner_system_message, quality_system_message, layman_system_message, planner_system_message
import copy
import pprint
import re
from typing import Dict, List, Tuple

import autogen
from autogen.agentchat.contrib.capabilities import transform_messages, transforms


# -------- CODEX

executor = LocalCommandLineCodeExecutor(
    timeout=3600,
    work_dir="./src/codex",
)


mardown_file_writer_agent = ConversableAgent(
    name="file_writer_agent",
    system_message = file_writer_agent_system_message,
    llm_config=llm_config,
    code_execution_config={"executor": executor},
    human_input_mode="NEVER",
    default_auto_reply=
    "Please continue. If everything is done, reply 'TERMINATE'.",
)


# --------  Tools

@file_writer_agent.register_for_execution()
@file_writer_agent.register_for_llm(description="Writes a text string to a markdown file")
def write_text_to_markdown(text, file_name, directory='./src/result/intermediate_results'):
    """
    Writes a text string to a markdown file, ensuring the directory exists.

    Args:
        text (str): The raw text content to be written to the markdown file.
        file_name (str): The name of the markdown file.
        directory (str): The directory where the markdown file will be saved. Default is './output_folder'.
    """
    # Ensure the directory exists
    os.makedirs(directory, exist_ok=True)
    
    # Full path to the markdown file
    file_path = os.path.join(directory, file_name)
    
    # Write the text content to the markdown file
    with open(file_path, 'w', encoding='utf-8') as file:
        file.write(text)
    
    print(f'Text has been written to {file_path}')

@file_writer_agent.register_for_execution()
@file_writer_agent.register_for_llm(description="Parses a markdown string into several markdown strings divided by titles, then saves each text string as a separate markdown document in a folder.")
def parse_markdown(markdown_str, output_folder="./src/result/intermediate_results"):
    """
    Parses a markdown string into several markdown strings divided by titles,
    then saves each text string as a separate markdown document in a folder.

    Args:
        markdown_str (str): The markdown content as a string.
        output_folder (str): The folder to save the numbered markdown files.

    Returns:
        List of filenames of the created markdown documents.
    """
    # Regex pattern to match markdown titles (assuming titles start with #)
    title_pattern = re.compile(r'(#+\s.*\n)')
    
    # Find all titles and their start positions
    matches = list(title_pattern.finditer(markdown_str))
    
    # Ensure the output folder exists
    os.makedirs(output_folder, exist_ok=True)
    
    filenames = []
    
    # Iterate over the matches and extract content
    for i in range(len(matches)):
        start_pos = matches[i].start()
        end_pos = matches[i + 1].start() if i + 1 < len(matches) else len(markdown_str)
        
        section_str = markdown_str[start_pos:end_pos]
        
        # Write the section to a separate markdown file
        filename = os.path.join(output_folder, f'section_{i + 1}.md')
        with open(filename, 'w', encoding='utf-8') as f:
            f.write(section_str)
        
        filenames.append(filename)
    
    return filenames

@file_writer_agent.register_for_execution()
@file_writer_agent.register_for_llm(description="Reads a markdown file and returns its raw text content.")
def read_markdown_file_to_text(file_path):
    """
    Reads a markdown file and returns its raw text content.

    Args:
        file_path (str): The path to the markdown file.

    Returns:
        str: The raw text content extracted from the markdown file.
    """
    # Read the markdown file
    with open(file_path, 'r', encoding='utf-8') as file:
        markdown_string = file.read()

    # Convert markdown to HTML
    html = markdown.markdown(markdown_string)
    
    # Use BeautifulSoup to extract text from HTML
    soup = BeautifulSoup(html, 'html.parser')
    text = soup.get_text()
    
    return text

@file_writer_agent.register_for_execution()
@file_writer_agent.register_for_llm(description="Appends text to the end of a given title in a .docx file (before the next title) and saves the updated document.")
def append_text_to_title( title: str, text_to_append: str, doc_path: str= '.src/result/result.docx', save_path: Optional[str] = None ) -> None:
    """
    Converts markdown to .docx and appends text to the end of a given title in a .docx file (before the next title) and saves the updated document.

    :param title: The title to search for in the document.
    :param text_to_append: The text to append to the end of the section with the specified title.
    :return: None
    """
    try:
        # Convert markdown text to Docx-suitable format
        formatted_text = markdown_to_docx_format(text_to_append)
     
    try:
        # Load the Document
        doc = Document(doc_path)
        
        # Iterate through paragraphs to find the title
        found_title = False
        insertion_point = None

        for i, para in enumerate(doc.paragraphs):
            if found_title:
                if para.style.name.startswith('Heading'):
                    # We have reached the next title
                    insertion_point = i
                    break
            elif para.text.strip() == title:
                found_title = True

        # If the title is found, insert the text
        if found_title:
            if insertion_point is None:
                # Append to the end of the document if no next title is found
                doc.add_paragraph(formatted_text)
            else:
                # Insert before the next title
                doc.paragraphs.insert(insertion_point, Paragraph(text_to_append, doc))
        
        # Save the updated document
        if save_path is None:
            save_path = doc_path
        doc.save(save_path)
        print(f"Document updated and saved at: {save_path}")

    except Exception as e:
        print(f"An error occurred while processing the .docx file: {e}")

# # Example usage
# doc_path = ".src/result/result.docx"
# title_to_search = "Introduction"
# text_to_append = "This is the appended text."
# append_text_to_title(doc_path, title_to_search, text_to_append)

@file_writer_agent.register_for_execution()
@file_writer_agent.register_for_llm(description="Replaces the text between the specified title and the next title in a .docx file with new text.")
def replace_section_with_text(title: str, new_text: str, doc_path: str= '.src/result/result.docx',  save_path: Optional[str] = None) -> None:
    """
    Converts Markdown text and Replaces the text between the specified title and the next title in a .docx file with new text.

    :param doc_path: The path to the .docx file to be read.
    :param title: The title that marks the beginning of the section to replace.
    :param new_text: The new text to insert in place of the specified section.
    :param save_path: The path to save the updated .docx file. If None, overwrites the original file.
    :return: None
    """
    try:
        # Convert markdown text to Docx-suitable format
        formatted_text = markdown_to_docx_format(new_text)
     
    try:
        # Load the Document
        doc = Document(doc_path)
        
        # Initialize variables to track the section boundaries
        section_start = None
        section_end = None

        # Find the title and the boundaries of the section to replace
        for i, para in enumerate(doc.paragraphs):
            if section_start is not None and para.style.name.startswith('Heading'):
                section_end = i
                break
            if para.text.strip() == title:
                section_start = i

        # If the title is found, replace the section
        if section_start is not None:
            # Remove old section paragraphs
            if section_end is None:
                # If no next title, clear till the end of the document
                section_end = len(doc.paragraphs)
            for _ in range(section_end - section_start - 1):
                del doc.paragraphs[section_start + 1]
                
            # Insert new text
            new_paragraphs = formatted_text.split("\n")
            for new_para in new_paragraphs:
                doc.paragraphs[section_start].insert_paragraph_before(new_para)
            del doc.paragraphs[section_start]  # Remove the old title paragraph as the new ones include the title

        # Save the updated document
        if save_path is None:
            save_path = doc_path
        doc.save(save_path)
        print(f"Document updated and saved at: {save_path}")

    except Exception as e:
        print(f"An error occurred while processing the .docx file: {e}")

def replace_specific_text_in_markdown(file_name: str, old_text: str, new_text: str , markdown_directory: str = "./src/result/intermediate_results"):
    """
    Opens a markdown file from a directory, replaces specified text, and saves/overrides the file.

    Args:
        markdown_directory (str): The directory containing the markdown file.
        file_name (str): The name of the markdown file to be modified.
        old_text (str): The text to search for and replace in the markdown file.
        new_text (str): The text to replace the old text with in the markdown file.
    
    Returns:
        None
    """
    # Construct the full file path
    file_path = os.path.join(markdown_directory, file_name)
    
    # Ensure the file exists in the specified directory
    if not os.path.isfile(file_path):
        raise ValueError(f"The specified markdown file does not exist: {file_path}")

    # Read the content of the markdown file
    with open(file_path, 'r', encoding='utf-8') as file:
        content = file.read()

    # Replace the specified text
    updated_content = content.replace(old_text, new_text)

    # Save the updated content back to the same file, overriding it
    with open(file_path, 'w', encoding='utf-8') as file:
        file.write(updated_content)

    print(f"Replaced text in {file_path} and saved the updated file.")

def replace_all_text_in_markdown(markdown_directory: str, file_name: str, new_text: str):
    """
    Opens a markdown file from a directory, replaces all text with new text, and saves/overrides the file.

    Args:
        markdown_directory (str): The directory containing the markdown file.
        file_name (str): The name of the markdown file to be modified.
        new_text (str): The new text to replace the existing content in the markdown file.
    
    Returns:
        None
    """
    # Construct the full file path
    file_path = os.path.join(markdown_directory, file_name)
    
    # Ensure the file exists in the specified directory
    if not os.path.isfile(file_path):
        raise ValueError(f"The specified markdown file does not exist: {file_path}")

    # Save the new content to the same file, overriding it
    with open(file_path, 'w', encoding='utf-8') as file:
        file.write(new_text)

    print(f"Replaced all text in {file_path} and saved the updated file.")

def replace_section_in_markdown(file_name: str, title: str, new_text: str, markdown_directory: str = ".src/result/intermediate_results"):
    """
    Opens a markdown file from a directory, finds a given title, and replaces all text
    until the next title or subtitle with the text provided, then saves/overrides the file.

    Args:
        markdown_directory (str): The directory containing the markdown file.
        file_name (str): The name of the markdown file to be modified.
        title (str): The title of the section to replace.
        new_text (str): The new text to replace the existing content in the section.
    
    Returns:
        None
    """
    # Construct the full file path
    file_path = os.path.join(markdown_directory, file_name)
    
    # Ensure the file exists in the specified directory
    if not os.path.isfile(file_path):
        raise ValueError(f"The specified markdown file does not exist: {file_path}")
    
    # Read the content of the markdown file
    with open(file_path, 'r', encoding='utf-8') as file:
        content = file.read()

    # Create a regex pattern to find the title and the subsequent content
    title_pattern = re.compile(r'^(#+\s+' + re.escape(title) + r'\s*\n)', re.MULTILINE)
    match = title_pattern.search(content)
    
    if not match:
        raise ValueError(f"Title '{title}' not found in the markdown file.")
    
    # Find the start and end positions of the section to replace
    start_pos = match.start()
    end_pos = len(content)
    
    next_title_pattern = re.compile(r'^#+\s+', re.MULTILINE)
    next_match = next_title_pattern.search(content, pos=match.end())
    
    if next_match:
        end_pos = next_match.start()
    
    # Create the new content by replacing the old section with the new text
    new_content = content[:match.end()] + new_text + content[end_pos:]
    
    # Save the updated content back to the same file, overriding it
    with open(file_path, 'w', encoding='utf-8') as file:
        file.write(new_content)

    print(f"Section under title '{title}' replaced and file saved at {file_path}.")
