from autogen import AssistantAgent, ConversableAgent, UserProxyAgent
from autogen.coding import LocalCommandLineCodeExecutor
from autogen.cache import Cache
from autogen.agentchat.contrib.math_user_proxy_agent  import MathUserProxyAgent
from src.config import llm_config
from docx import Document
from docx.text.paragraph import Paragraph
import os
from typing import Optional, List, Tuple
from markdown2 import markdown
from bs4 import BeautifulSoup
from src.prompts import file_writer_agent_system_message, docx_planner_system_message
import copy
import pprint
import re
from typing import Dict, List, Tuple

import autogen
from autogen.agentchat.contrib.capabilities import transform_messages, transforms


# --------  Docx Planner



docx_planner = AssistantAgent(
    name="docx file writer planner",
    llm_config=llm_config,
    # the default system message of the AssistantAgent is overwritten here
    system_message=docx_planner_system_message,
)

docx_planner_user = UserProxyAgent(
    name="docx_planner_user",
    max_consecutive_auto_reply=0,  # terminate without auto-reply
    human_input_mode="NEVER",
    code_execution_config={
        "use_docker": True
    },  # Please set use_docker=True if docker is available to run the generated code. Using docker is safer than running the generated code directly.
)

def docx_ask_planner(message):
    docx_planner_user.initiate_chat(docx_planner, message=message)
    # return the last message received from the planner
    return docx_planner_user.last_message()["content"]

docx_assistant = AssistantAgent(
    name="assistant",
    llm_config={
        "temperature": 0,
        "timeout": 600,
        "cache_seed": 42,
        "config_list": llm_config,
        "functions": [
            {
                "name": "docx_ask_planner",
                "description": "ask docx_planner to: 1. get a plan for finishing a task, 2. verify the execution result of the plan and potentially suggest new plan.",
                "parameters": {
                    "type": "object",
                    "properties": {
                        "message": {
                            "type": "string",
                            "description": "question to ask planner. Make sure the question include enough context, such as the code and the execution result. The planner does not know the conversation between you and the user, unless you share the conversation with the planner.",
                        },
                    },
                    "required": ["message"],
                },
            },
        ],
    },
)

# create a UserProxyAgent instance named "user_proxy"
docx_user_proxy = UserProxyAgent(
    name="docx_user_proxy",
    human_input_mode="TERMINATE",
    max_consecutive_auto_reply=10,
    # is_termination_msg=lambda x: "content" in x and x["content"] is not None and x["content"].rstrip().endswith("TERMINATE"),
    code_execution_config={
        "work_dir": "./src/codex",
        "use_docker": False,
    },  # Please set use_docker=True if docker is available to run the generated code. Using docker is safer than running the generated code directly.
    function_map={"docx_ask_planner": docx_ask_planner},
)


# -------- CODEX

docx_executor = LocalCommandLineCodeExecutor(
    timeout=3600,
    work_dir="./src/codex",
)


docx_file_writer_agent = ConversableAgent(
    name="file_writer_agent",
    system_message = file_writer_agent_system_message,
    llm_config=llm_config,
    code_execution_config={"executor": docx_executor},
    human_input_mode="NEVER",
    default_auto_reply=
    "Please continue. If everything is done, reply 'TERMINATE'.",
)

# --------  Tools

@docx_file_writer_agent.register_for_execution()
@docx_file_writer_agent.register_for_llm(description="Reads and loads the data from a .docx file as plain text and returns the complete file content as plain text.")
def read_docx_as_plain_text(filepath: str = '.src/result/result.docx') -> Optional[str]:
    """
    Reads and loads the data from a .docx file as plain text and returns the complete file content as plain text.

    :param filepath: The path to the .docx file to be read.
    :return: The plain text content of the .docx file, or None if the file cannot be read.
    """
    try:
        # Load the Document
        doc = Document(filepath)
        
        # Initialize an empty string to hold the text
        full_text = []
        
        # Iterate through each paragraph in the document
        for para in doc.paragraphs:
            full_text.append(para.text)
        
        # Join all paragraphs into a single string separated by newlines
        result_text = "\n".join(full_text)
        
        return result_text
    except Exception as e:
        print(f"An error occurred while reading the .docx file: {e}")
        return None
    
@docx_file_writer_agent.register_for_execution()
@docx_file_writer_agent.register_for_llm(description="Appends text to the end of a given title in a .docx file (before the next title) and saves the updated document.")
def append_text_to_title( title: str, text_to_append: str, doc_path: str= '.src/result/result.docx', save_path: Optional[str] = None ) -> None:
    """
    Converts markdown to .docx and appends text to the end of a given title in a .docx file (before the next title) and saves the updated document.

    :param title: The title to search for in the document.
    :param text_to_append: The text to append to the end of the section with the specified title.
    :return: None
    """
    try:
        # Convert markdown text to Docx-suitable format
        formatted_text = markdown_to_docx_format(text_to_append)

    except Exception as e:
        print(f"An error occurred while converting markdown to docx format: {e}")
        return ""
             
    try:
        # Load the Document
        doc = Document(doc_path)
        
        # Iterate through paragraphs to find the title
        found_title = False
        insertion_point = None

        for i, para in enumerate(doc.paragraphs):
            if found_title:
                if para.style.name.startswith('Heading'):
                    # We have reached the next title
                    insertion_point = i
                    break
            elif para.text.strip() == title:
                found_title = True

        # If the title is found, insert the text
        if found_title:
            if insertion_point is None:
                # Append to the end of the document if no next title is found
                doc.add_paragraph(formatted_text)
            else:
                # Insert before the next title
                doc.paragraphs.insert(insertion_point, Paragraph(text_to_append, doc))
        
        # Save the updated document
        if save_path is None:
            save_path = doc_path
        doc.save(save_path)
        print(f"Document updated and saved at: {save_path}")

    except Exception as e:
        print(f"An error occurred while processing the .docx file: {e}")

# # Example usage
# doc_path = ".src/result/result.docx"
# title_to_search = "Introduction"
# text_to_append = "This is the appended text."
# append_text_to_title(doc_path, title_to_search, text_to_append)

@docx_file_writer_agent.register_for_execution()
@docx_file_writer_agent.register_for_llm(description="Replaces the text between the specified title and the next title in a .docx file with new text.")
def replace_section_with_text(title: str, new_text: str, doc_path: str= '.src/result/result.docx',  save_path: Optional[str] = None) -> None:
    """
    Converts Markdown text and Replaces the text between the specified title and the next title in a .docx file with new text.

    :param doc_path: The path to the .docx file to be read.
    :param title: The title that marks the beginning of the section to replace.
    :param new_text: The new text to insert in place of the specified section.
    :param save_path: The path to save the updated .docx file. If None, overwrites the original file.
    :return: None
    """
    try:
        # Convert markdown text to Docx-suitable format
        formatted_text = markdown_to_docx_format(new_text)
    
    except Exception as e:
        print(f"An error occurred while converting markdown to docx format: {e}")
        return "" 
            
    try:
        # Load the Document
        doc = Document(doc_path)
        
        # Initialize variables to track the section boundaries
        section_start = None
        section_end = None

        # Find the title and the boundaries of the section to replace
        for i, para in enumerate(doc.paragraphs):
            if section_start is not None and para.style.name.startswith('Heading'):
                section_end = i
                break
            if para.text.strip() == title:
                section_start = i

        # If the title is found, replace the section
        if section_start is not None:
            # Remove old section paragraphs
            if section_end is None:
                # If no next title, clear till the end of the document
                section_end = len(doc.paragraphs)
            for _ in range(section_end - section_start - 1):
                del doc.paragraphs[section_start + 1]
                
            # Insert new text
            new_paragraphs = formatted_text.split("\n")
            for new_para in new_paragraphs:
                doc.paragraphs[section_start].insert_paragraph_before(new_para)
            del doc.paragraphs[section_start]  # Remove the old title paragraph as the new ones include the title

        # Save the updated document
        if save_path is None:
            save_path = doc_path
        doc.save(save_path)
        print(f"Document updated and saved at: {save_path}")

    except Exception as e:
        print(f"An error occurred while processing the .docx file: {e}")

@docx_file_writer_agent.register_for_execution()
@docx_file_writer_agent.register_for_llm(description="Converts markdown content to a single string with python-docx compatible formatting.")
def markdown_to_docx_format(markdown_text: str) -> str:
    """
    Converts markdown content to a single string with python-docx compatible formatting.

    :param markdown_text: The raw text containing markdown elements.
    :return: A single string with the complete text, formatted for further processing.
    """
    try:
        # Convert markdown to HTML
        html = markdown(markdown_text)
        
        # Parse HTML using BeautifulSoup
        soup = BeautifulSoup(html, 'html.parser')
        
        formatted_text = ""

        # Define mapping for header tags
        header_map = {
            'h1': ('Heading 1', 1),
            'h2': ('Heading 2', 2),
            'h3': ('Heading 3', 3),
            'h4': ('Heading 4', 4),
            'h5': ('Heading 5', 5),
            'h6': ('Heading 6', 6)
        }
        
        # Helper function to add text with possible styling
        def add_styled_text(element):
            text = ''
            if element.name == 'strong':
                text += '**' + element.text + '**'
            elif element.name == 'em':
                text += '*' + element.text + '*'
            elif element.name == 'code':
                text += '`' + element.text + '`'
            else:
                text += element.text
            return text

        # Iterate through HTML elements
        for tag in soup.contents:
            if tag.name in header_map:
                formatted_text += f"\n\n{header_map[tag.name][0]}: {tag.text.strip()}\n"
            elif tag.name == 'p':
                p_text = ''.join(add_styled_text(child) for child in tag.children)
                formatted_text += f"\n{p_text}\n"
            elif tag.name == 'ul':
                for li in tag.find_all('li'):
                    formatted_text += f"\n- {li.text.strip()}"
            elif tag.name == 'ol':
                for li in tag.find_all('li'):
                    formatted_text += f"\n1. {li.text.strip()}"
            elif tag.name == 'pre':
                code_block = tag.find('code')
                if code_block:
                    formatted_text += f"\n\n```\n{code_block.text}\n```\n"

        return formatted_text.strip()
    
    except Exception as e:
        print(f"An error occurred while converting markdown to docx format: {e}")
        return ""


@docx_file_writer_agent.register_for_execution()
@docx_file_writer_agent.register_for_llm(description="Converts markdown text to docx format and writes it to a .docx file.")
def write_markdown_to_docx(markdown_text: str, save_path: Optional[str] = "./src/result/result.docx", doc_path: str = "./src/result/result.docx", ) -> None:
    """
    Converts markdown text to docx format and writes it to a .docx file.

    :param markdown_text: The raw text containing markdown elements.
    :param doc_path: The path to the existing .docx file.
    :param save_path: The path to save the updated .docx file. If None, overwrites the original file.
    :return: None
    """
    try:
        # Convert markdown text to Docx-suitable format
        formatted_text = markdown_to_docx_format(markdown_text)
        
        # Load the Document
        doc = Document(doc_path)
        
        # Split the formatted text into paragraphs based on newline separators
        paragraphs = formatted_text.split('\n\n')
        
        for paragraph in paragraphs:
            if ': ' in paragraph:
                # Check for headers
                header_type, text = paragraph.split(': ', 1)
                doc.add_heading(text.strip(), level=header_type[-1])
            elif paragraph.startswith('```') and paragraph.endswith('```'):
                # Check for code blocks
                code_text = paragraph.strip('```').strip()
                doc.add_paragraph(code_text, style='Quote')
            else:
                # Handle bullet points and other text
                lines = paragraph.split('\n')
                for line in lines:
                    if line.startswith('- '):
                        doc.add_paragraph(line.strip('- '), style='List Bullet')
                    elif line.startswith('1. '):
                        doc.add_paragraph(line.strip('1. '), style='List Number')
                    else:
                        doc.add_paragraph(line)

        # Save the updated document
        if save_path is None:
            save_path = doc_path
        doc.save(save_path)
        print(f"Document updated and saved at: {save_path}")

    except Exception as e:
        print(f"An error occurred while converting markdown to docx format: {e}")
        return ""    


@docx_file_writer_agent.register_for_execution()
@docx_file_writer_agent.register_for_llm(description="Iterates over each markdown file in a directory and writes them to a .docx file.")
def process_markdown_files_in_directory(markdown_directory: str = "./src/result/intermediate_results", docx_save_path: str = "./src/result/result.docx"):
    """
    Iterates over each markdown file in a directory and writes them to a .docx file.

    Args:
        markdown_directory (str): The directory containing markdown files.
        docx_save_path (str): The path to save the combined .docx file.

    Returns:
        None
    """
    # Ensure the markdown directory exists
    if not os.path.isdir(markdown_directory):
        raise ValueError(f"The specified directory does not exist: {markdown_directory}")
    
    # Create a new Document or load an existing one
    if os.path.exists(docx_save_path):
        doc = Document(docx_save_path)
    else:
        doc = Document()
    
    # Iterate over each markdown file in the directory
    for filename in os.listdir(markdown_directory):
        if filename.endswith('.md'):
            filepath = os.path.join(markdown_directory, filename)
            
            # Read the markdown file content
            with open(filepath, 'r', encoding='utf-8') as file:
                markdown_text = file.read()
            
            # Convert markdown text to docx format and append it to the document
            formatted_text = markdown_to_docx_format(markdown_text)
            paragraphs = formatted_text.split('\n\n')
            
            for paragraph in paragraphs:
                if ': ' in paragraph:
                    # Check for headers
                    header_type, text = paragraph.split(': ', 1)
                    doc.add_heading(text.strip(), level=int(header_type[-1]))
                elif paragraph.startswith('```') and paragraph.endswith('```'):
                    # Check for code blocks
                    code_text = paragraph.strip('```').strip()
                    doc.add_paragraph(code_text, style='Quote')
                else:
                    # Handle bullet points and other text
                    lines = paragraph.split('\n')
                    for line in lines:
                        if line.startswith('- '):
                            doc.add_paragraph(line.strip('- '), style='List Bullet')
                        elif line.startswith('1. '):
                            doc.add_paragraph(line.strip('1. '), style='List Number')
                        else:
                            doc.add_paragraph(line)

    # Save the combined document
    doc.save(docx_save_path)
    print(f"Combined document saved at: {docx_save_path}")
