# /src/utils.py Tools

from markdown2 import markdown
from bs4 import BeautifulSoup
from docx import Document  
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_LINE_SPACING, WD_BREAK  
from docx.shared import Pt, Inches  
from docx.oxml.ns import qn  
from docx.oxml import OxmlElement  
from docx.enum.style import WD_STYLE_TYPE  
from docx.text.paragraph import Paragraph
from typing import Optional,  Dict, List, Tuple
import re  
import os  
import json  
import logging  
import glob  

def write_text_to_markdown(text, file_name, directory='./src/result/intermediate_results'):
    """
    Writes a text string to a markdown file, ensuring the directory exists.

    Args:
        text (str): The raw text content to be written to the markdown file.
        file_name (str): The name of the markdown file.
        directory (str): The directory where the markdown file will be saved. Default is './output_folder'.
    """
    # Ensure the directory exists
    os.makedirs(directory, exist_ok=True)
    
      
    # Ensure the file_name does not contain any path segments  
    file_name = os.path.basename(file_name)  
  
    # Full path to the markdown file
    file_path = os.path.join(directory, file_name)

    # Convert text to string if it's a dictionary or list  
    if isinstance(text, dict):  
        text = json.dumps(text, indent=4)  # Pretty-print the JSON with indentation  
    elif isinstance(text, list):  
        text = '\n'.join(text)  # Convert list to a string with each element on a new line  
  
    # Write the text content to the markdown file
    with open(file_path, 'w', encoding='utf-8') as file:
        file.write(text)
    
    print(f'Text has been written to {file_path}')

def parse_markdown(markdown_str, output_folder="./src/result/intermediate_results"):  
    """Parses a markdown string into several markdown strings divided by titles, then saves each text string as a separate markdown document in a folder."""  
    title_pattern = re.compile(r'(#+\s.*\n)')  
    matches = list(title_pattern.finditer(markdown_str))  
    os.makedirs(output_folder, exist_ok=True)  
    filenames = []  
    for i in range(len(matches)):  
        start_pos = matches[i].start()  
        end_pos = matches[i + 1].start() if i + 1 < len(matches) else len(markdown_str)  
        section_str = markdown_str[start_pos:end_pos]  
        filename = os.path.join(output_folder, f'section_{i + 1}.md')  
        with open(filename, 'w', encoding='utf-8') as f:  
            f.write(section_str)  
        filenames.append(filename)  
    return filenames

def read_markdown_file_to_text(file_path):
    """
    Reads a markdown file and returns its raw text content.

    Args:
        file_path (str): The path to the markdown file.

    Returns:
        str: The raw text content extracted from the markdown file.
    """
    # Read the markdown file
    with open(file_path, 'r', encoding='utf-8') as file:
        markdown_string = file.read()

    # Convert markdown to HTML
    html = markdown(markdown_string)
    
    # Use BeautifulSoup to extract text from HTML
    soup = BeautifulSoup(html, 'html.parser')
    text = soup.get_text()
    
    return text



def clear_previous_results(intermediate_results_directory, final_result_path):  
    """Remove all files in the intermediate results directory and the final result file if they exist."""  
    try:  
        files = glob.glob(f"{intermediate_results_directory}/*")  
        for f in files:  
            os.remove(f)  
        if os.path.exists(final_result_path):  
            os.remove(final_result_path)  
        logging.info(f"Cleared previous results in {intermediate_results_directory} and removed {final_result_path} if it existed.")  
    except Exception as e:  
        logging.error(f"Error clearing previous results: {e}")  


def markdown_to_docx_format(markdown_text: str, doc: Document) -> None:  
    """Converts markdown content to a python-docx document."""  
    try:  
        header_map = {  
            'h1': 0,  
            'h2': 1,  
            'h3': 2,  
            'h4': 3,  
            'h5': 4,  
            'h6': 5  
        }  
  
        def add_styled_text(paragraph, text):  
            """Add styled text (bold, italic, code) to a paragraph."""  
            parts = re.split(r'(\*\*.*?\*\*|\*.*?\*|`.*?`)', text)  
            for part in parts:  
                if part.startswith('**') and part.endswith('**'):  
                    run = paragraph.add_run(part[2:-2])  
                    run.bold = True  
                elif part.startswith('*') and part.endswith('*'):  
                    run = paragraph.add_run(part[1:-1])  
                    run.italic = True  
                elif part.startswith('`') and part.endswith('`'):  
                    run = paragraph.add_run(part[1:-1])  
                    run.font.name = 'Courier New'  
                else:  
                    paragraph.add_run(part)  
  
        list_stack = []  
        current_list_type = None  
        lines = markdown_text.split('\n')  
  
        for line in lines:  
            if re.match(r'#+\s', line):  
                # Check for headers  
                match = re.match(r'^(#+)\s+(.*)', line)  
                if match:  
                    header_level = len(match.group(1))  
                    header_text = match.group(2)  
                    doc.add_heading(header_text.strip(), level=header_map.get(f'h{header_level}', 1))  
            elif line.startswith('```') and line.endswith('```'):  
                # Check for code blocks  
                code_text = line.strip('```').strip()  
                doc.add_paragraph(code_text, style='Quote')  
            elif line.startswith('- '):  
                # Handle unordered list  
                if current_list_type != 'ul':  
                    current_list_type = 'ul'  
                    list_stack.append(current_list_type)  
                paragraph = doc.add_paragraph(line[2:].strip(), style='List Bullet')  
            elif re.match(r'^\d+\.', line):  
                # Handle ordered list  
                if current_list_type != 'ol':  
                    current_list_type = 'ol'  
                    list_stack.append(current_list_type)  
                paragraph = doc.add_paragraph(line.strip(), style='List Number')  
            elif '|' in line and '-' not in line:  
                # Handle table rows  
                cells = [cell.strip() for cell in line.split('|') if cell.strip()]  
                if not hasattr(doc, '_current_table'):  
                    doc._current_table = doc.add_table(rows=1, cols=len(cells))  
                    doc._current_table.style = 'Table Grid'  
                    hdr_cells = doc._current_table.rows[0].cells  
                    for i, cell in enumerate(cells):  
                        hdr_cells[i].text = cell  
                else:  
                    row_cells = doc._current_table.add_row().cells  
                    for i, cell in enumerate(cells):  
                        row_cells[i].text = cell  
            else:  
                # Handle normal paragraphs  
                paragraph = doc.add_paragraph()  
                add_styled_text(paragraph, line.strip())  
  
        if hasattr(doc, '_current_table'):  
            del doc._current_table  
  
    except Exception as e:  
        print(f"An error occurred while converting markdown to docx format: {e}")  
  
def markdown_to_docx(markdown_directory: str = "./src/result/intermediate_results", docx_save_path: str = "./src/result/result.docx"):  
    """Iterates over each markdown file in a directory, converts them to a single string with python-docx compatible formatting, and writes them to a .docx file."""  
    if not os.path.isdir(markdown_directory):  
        raise ValueError(f"The specified directory does not exist: {markdown_directory}")  
  
    doc = Document()  
  
    for filename in os.listdir(markdown_directory):  
        if filename.endswith('.md'):  
            filepath = os.path.join(markdown_directory, filename)  
            with open(filepath, 'r', encoding='utf-8') as file:  
                markdown_text = file.read()  
            markdown_to_docx_format(markdown_text, doc)  
  
    doc.save(docx_save_path)  
    print(f"Combined document saved at: {docx_save_path}")  