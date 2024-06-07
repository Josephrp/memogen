# src/main.py

import autogen
from src.agents_writer import layman_reviewer, financial_reviewer, quality_reviewer , outliner,  meta_reviewer, critic, writer
# from src.agents_docx import  process_markdown_files_in_directory
# from src.agents_markdown import  parse_markdown, read_markdown_file_to_text, write_text_to_markdown
from src.config import load_env_file
# from autogen.agentchat.groupchat import GroupChat
from markdown2 import markdown
from bs4 import BeautifulSoup
from autogen import initiate_chats
from autogen.cache import Cache
from src.config import llm_config
from docx import Document
from docx.text.paragraph import Paragraph
from typing import Optional,  Dict, List, Tuple
import logging
import os
import re

# ---------- Init


# load_env_file('/.env')

# chroma_client = chromadb.HttpClient(host='localhost', port=8000)

topic = ""
audience = ""
memo_type = ""

# ----------- Tools

def write_text_to_markdown(text, file_name, directory='./src/result/intermediate_results'):
    """
    Writes a text string to a markdown file, ensuring the directory exists.

    Args:
        text (str): The raw text content to be written to the markdown file.
        file_name (str): The name of the markdown file.
        directory (str): The directory where the markdown file will be saved. Default is './output_folder'.
    """
    # Ensure the directory exists
    os.makedirs(directory, exist_ok=True)
    
    # Full path to the markdown file
    file_path = os.path.join(directory, file_name)
    
    # Write the text content to the markdown file
    with open(file_path, 'w', encoding='utf-8') as file:
        file.write(text)
    
    print(f'Text has been written to {file_path}')

def markdown_to_docx_format(markdown_text: str) -> str:
    """
    Converts markdown content to a single string with python-docx compatible formatting.

    :param markdown_text: The raw text containing markdown elements.
    :return: A single string with the complete text, formatted for further processing.
    """
    try:
        # Convert markdown to HTML
        html = markdown(markdown_text)

        # Parse HTML using BeautifulSoup
        soup = BeautifulSoup(html, 'html.parser')

        formatted_text = ""

        # Define mapping for header tags
        header_map = {
            'h1': ('Heading 1', 1),
            'h2': ('Heading 2', 2),
            'h3': ('Heading 3', 3),
            'h4': ('Heading 4', 4),
            'h5': ('Heading 5', 5),
            'h6': ('Heading 6', 6)
        }

        # Helper function to add text with possible styling
        def add_styled_text(element):
            text = ''
            if element.name == 'strong':
                text += '**' + element.text + '**'
            elif element.name == 'em':
                text += '*' + element.text + '*'
            elif element.name == 'code':
                text += '`' + element.text + '`'
            else:
                text += element.text
            return text

        # Iterate through HTML elements
        for tag in soup.contents:
            if tag.name in header_map:
                formatted_text += f"\n\n{header_map[tag.name][0]}: {tag.text.strip()}\n"
            elif tag.name == 'p':
                p_text = ''.join(add_styled_text(child) for child in tag.children)
                formatted_text += f"\n{p_text}\n"
            elif tag.name == 'ul':
                for li in tag.find_all('li'):
                    formatted_text += f"\n- {li.text.strip()}"
            elif tag.name == 'ol':
                for li in tag.find_all('li'):
                    formatted_text += f"\n1. {li.text.strip()}"
            elif tag.name == 'pre':
                code_block = tag.find('code')
                if code_block:
                    formatted_text += f"\n\n```\n{code_block.text}\n```\n"

        return formatted_text.strip()

    except Exception as e:
        print(f"An error occurred while converting markdown to docx format: {e}")
        return ""

def write_markdown_to_docx(markdown_text: str, save_path: Optional[str] = "./src/result/result.docx", doc_path: str = "./src/result/result.docx", ) -> None:
    """
    Converts markdown text to docx format and writes it to a .docx file.

    :param markdown_text: The raw text containing markdown elements.
    :param doc_path: The path to the existing .docx file.
    :param save_path: The path to save the updated .docx file. If None, overwrites the original file.
    :return: None
    """
    try:
        # Convert markdown text to Docx-suitable format
        formatted_text = markdown_to_docx_format(markdown_text)

        # Load the Document
        doc = Document(doc_path)

        # Split the formatted text into paragraphs based on newline separators
        paragraphs = formatted_text.split('\n\n')

        for paragraph in paragraphs:
            if ': ' in paragraph:
                # Check for headers
                header_type, text = paragraph.split(': ', 1)
                doc.add_heading(text.strip(), level=header_type[-1])
            elif paragraph.startswith('```') and paragraph.endswith('```'):
                # Check for code blocks
                code_text = paragraph.strip('```').strip()
                doc.add_paragraph(code_text, style='Quote')
            else:
                # Handle bullet points and other text
                lines = paragraph.split('\n')
                for line in lines:
                    if line.startswith('- '):
                        doc.add_paragraph(line.strip('- '), style='List Bullet')
                    elif line.startswith('1. '):
                        doc.add_paragraph(line.strip('1. '), style='List Number')
                    else:
                        doc.add_paragraph(line)

        # Save the updated document
        if save_path is None:
            save_path = doc_path
        doc.save(save_path)
        print(f"Document updated and saved at: {save_path}")

    except Exception as e:
        print(f"An error occurred while converting markdown to docx format: {e}")
        return ""

def process_markdown_files_in_directory(markdown_directory: str = "./src/result/intermediate_results", docx_save_path: str = "./src/result/result.docx"):
    """
    Iterates over each markdown file in a directory and writes them to a .docx file.

    Args:
        markdown_directory (str): The directory containing markdown files.
        docx_save_path (str): The path to save the combined .docx file.

    Returns:
        None
    """
    # Ensure the markdown directory exists
    if not os.path.isdir(markdown_directory):
        raise ValueError(f"The specified directory does not exist: {markdown_directory}")

    # Create a new Document or load an existing one
    if os.path.exists(docx_save_path):
        doc = Document(docx_save_path)
    else:
        doc = Document()

    # Iterate over each markdown file in the directory
    for filename in os.listdir(markdown_directory):
        if filename.endswith('.md'):
            filepath = os.path.join(markdown_directory, filename)

            # Read the markdown file content
            with open(filepath, 'r', encoding='utf-8') as file:
                markdown_text = file.read()

            # Convert markdown text to docx format and append it to the document
            formatted_text = markdown_to_docx_format(markdown_text)
            paragraphs = formatted_text.split('\n\n')

            for paragraph in paragraphs:
                if ': ' in paragraph:
                    # Check for headers
                    header_type, text = paragraph.split(': ', 1)
                    doc.add_heading(text.strip(), level=int(header_type[-1]))
                elif paragraph.startswith('```') and paragraph.endswith('```'):
                    # Check for code blocks
                    code_text = paragraph.strip('```').strip()
                    doc.add_paragraph(code_text, style='Quote')
                else:
                    # Handle bullet points and other text
                    lines = paragraph.split('\n')
                    for line in lines:
                        if line.startswith('- '):
                            doc.add_paragraph(line.strip('- '), style='List Bullet')
                        elif line.startswith('1. '):
                            doc.add_paragraph(line.strip('1. '), style='List Number')
                        else:
                            doc.add_paragraph(line)

    # Save the combined document
    doc.save(docx_save_path)
    print(f"Combined document saved at: {docx_save_path}")

# --------- User Input

def get_user_inputs():
    """Prompt and get inputs from the user for topic, audience, and type of memo with clear instructions."""
    print("Welcome to the Memo Production System. Please provide the following information:")

    topic = input("1. Enter the topic for the memo (e.g., Artificial Intelligence in Modern Healthcare, Accounting Memo To Explain the Benefits of Going Public): ")

    audience = input("2. Enter the audience for the memo (e.g., Healthcare Professionals and Administrators): ")

    memo_type = input("3. Enter the type of memo (e.g., Accounting, Financial, Technical, Policy...): ")

    print("\nThank you! Creating the memo based on the provided details...\n")

    return topic, audience, memo_type

# --------- Agents 

from autogen import AssistantAgent
from src.prompts import get_system_messages
from src.config import llm_config
import os

# Fetch all system messages using the appropriate function
messages = get_system_messages(audience="general", memo_type="General")
# -------- Writer

writer = AssistantAgent(
    name="Writer",
    system_message=messages["writer_system_message"],
    llm_config=llm_config,
)

outliner = AssistantAgent(
    name="Writer",
    system_message=messages["outliner_system_message"],
    llm_config=llm_config,
)

# -------- Reviewers

critic = AssistantAgent(
    name="Critic",
    is_termination_msg=lambda x: x.get("content", "").find("TERMINATE") >= 0,
    llm_config=llm_config,
    system_message=messages["critic_system_message"],
)

layman_reviewer = AssistantAgent(
    name="Layman Reviewer",
    description="A reviewer that makes sure a laywoman would fully understand the content provided to her.",
    llm_config=llm_config,
    system_message=messages["layman_system_message"],
)

financial_reviewer = AssistantAgent(
    name="Financial Reviewer",
    description='A reviewer that makes sure financial justification are credible',
    llm_config=llm_config,
    system_message=messages["financial_reviewer_system_message"],
)

quality_reviewer = AssistantAgent(
    name="Quality Assurance Reviewer",
    description="a reviewer that makes sure that claims are well justified",
    llm_config=llm_config,
    system_message=messages["quality_system_message"],
)

meta_reviewer = AssistantAgent(
    name="Meta Reviewer",
    llm_config=llm_config,
    system_message="You are a meta reviewer, you aggragate and review "
    "the work of other reviewers and give a final suggestion on the content.",
)


# --------- Main Application Logic

def reflection_message(recipient, messages, sender, config):
    return f'''Review the following content.
            \n\n {recipient.chat_messages_for_summary(sender)[-1]['content']}'''

review_chats = [
    {
     "recipient": layman_reviewer,
     "message": reflection_message,
     "summary_method": "reflection_with_llm",
     "summary_args": {"summary_prompt" :
        "Return review into as JSON object only:"
        "{'Reviewer': '', 'Review': ''}. Here Reviewer should be your role",},
     "max_turns": 1},
    {
    "recipient": financial_reviewer, "message": reflection_message,
     "summary_method": "reflection_with_llm",
     "summary_args": {"summary_prompt" :
        "Return review into as JSON object only:"
        "{'Reviewer': '', 'Review': ''}.",},
     "max_turns": 1},
    {"recipient": quality_reviewer, "message": reflection_message,
     "summary_method": "reflection_with_llm",
     "summary_args": {"summary_prompt" :
        "Return review into as JSON object only:"
        "{'reviewer': '', 'review': ''}",},
     "max_turns": 1},
     {"recipient": meta_reviewer,
      "message": "Aggregrate feedback from all reviewers and give final suggestions on the writing.",
     "max_turns": 1},
]

critic.register_nested_chats(
    review_chats,
    trigger=writer,
)


class AutoMemoProduction:
    def __init__(self, topic, audience, memo_type):
        logging.info(f"Initializing AgentManager")
        self.topic = topic
        self.audience = audience
        self.memo_type = memo_type

    def create_outline(self):
        message=f"Create an outline for a {self.memo_type} memo on the topic {self.topic} optimized for the audience: {self.audience}."
        outliner.generate_reply(messages=[{"content": message, "role": "user"}])
        outline = outliner.last_message().get("content", "")
        logging.info(f"Outline created: {outline}")
        return outline

    def parse_outline_to_markdown_chunks(self, outline_str):
        filenames = parse_markdown(outline_str)
        logging.info(f"Parsed markdown files: {filenames}")
        return filenames

    def write_sections(self, markdown_filenames):
        for filename in markdown_filenames:
            file_content = read_markdown_file_to_text(filename)
            # first_draft = writer.generate_reply(messages=[{"content": file_content, "role": "user"}])
            result = writer.initiate_chat(review_chats, message=f"{file_content} \n \n Produce a detailed section based on the section of {self.memo_type} memo on the topic of {self.topic} optimized for {self.audience} provided above:",)
            write_text_to_markdown(result, filename)
            logging.info(f"Completed writing section for {filename}")

    def combine_sections_to_docx(self, markdown_directory, docx_path):
        process_markdown_files_in_directory(markdown_directory, docx_path)
        logging.info(f"Combined markdown sections into docx: {docx_path}")

    def run(self):
        outline = self.create_outline()
        markdown_filenames = self.parse_outline_to_markdown_chunks(outline)
        self.write_sections(markdown_filenames)
        self.combine_sections_to_docx("./src/result/intermediate_results", "./src/result/final_memo.docx")


if __name__ == "__main__":
    logging.basicConfig(level=logging.INFO)
    topic, audience, memo_type = get_user_inputs()
    producer = AutoMemoProduction(topic=topic, audience=audience, memo_type=memo_type)
    producer.run()


